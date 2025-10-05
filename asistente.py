import streamlit as st
import requests
from bs4 import BeautifulSoup
import time
from datetime import datetime
import google.generativeai as genai
from google.generativeai import types
import io
import csv
from docx import Document
from semantic_scholar import SemanticScholar
import re

# Clase Agente con memoria
class AgenteConMemoria:
    def __init__(self):
        self.memoria = []
    
    def agregar_memoria(self, entrada, salida):
        self.memoria.append({"entrada": entrada, "salida": salida, "timestamp": datetime.now()})
    
    def obtener_memoria(self):
        return self.memoria
    
    def limpiar_memoria(self):
        self.memoria = []

# FUNCIONES DE BÚSQUEDA MEJORADAS
def buscar_semantic_scholar(query, max_results=5):
    """Búsqueda gratuita en Semantic Scholar"""
    articulos = []
    try:
        scholar = SemanticScholar()
        resultados = scholar.search_paper(query, limit=max_results)
        count = 0
        for paper in resultados:
            articulos.append({
                "titulo": paper.title,
                "autor": ". ".join([author['name'] for author in paper.authors]) if paper.authors else 'Desconocido',
                "año": str(paper.year) if paper.year else "s.f.",
                "publicacion": paper.venue if paper.venue else "",
                "url": paper.url if paper.url else "",
                "fuente": "Semantic Scholar"
            })
            count += 1
            if count >= max_results:
                break
        time.sleep(0.5)
    except Exception as e:
        st.error(f"Error en Semantic Scholar: {e}")
    return articulos

def buscar_scielo(query, max_results=5):
    """Búsqueda gratuita en SciELO"""
    articulos = []
    try:
        url = f"https://search.scielo.org/?q={query}&lang=es"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=30)
        soup = BeautifulSoup(response.text, 'html.parser')
        resultados = soup.select(".item")
        
        for resultado in resultados[:max_results]:
            try:
                titulo_elem = resultado.select_one(".title")
                autores_elem = resultado.select_one(".authors")
                enlace_elem = resultado.select_one(".line a")
                
                titulo = titulo_elem.text.strip() if titulo_elem else "Sin titulo"
                autores = autores_elem.text.strip() if autores_elem else "Desconocido"
                enlace = enlace_elem["href"] if enlace_elem else ""
                
                articulos.append({
                    "titulo": titulo,
                    "autor": autores,
                    "año": "s.f.",
                    "publicacion": "SciELO",
                    "url": f"https:{enlace}" if enlace else "",
                    "fuente": "SciELO"
                })
            except Exception:
                continue
    except Exception as e:
        st.error(f"Error en SciELO: {e}")
    return articulos

# FUNCIONES DE PROCESAMIENTO DE LENGUAJE MEJORADAS
def extraer_tema_principal(user_input):
    """Extrae el tema real de investigación del input del usuario"""
    try:
        # Patrón para detectar preguntas de investigación
        patron_pregunta = r'[¿]?(qué|cuales|como|por qué|dónde|cuándo)\s+([^?]+)[?]'
        coincidencia = re.search(patron_pregunta, user_input.lower())
        
        if coincidencia:
            return coincidencia.group(2).strip()
        
        # Eliminar palabras de solicitud metodológica
        palabras_excluir = ['formula', 'planteamiento', 'problema', 'interrogante', 
                           'redacta', 'elabora', 'desarrolla', 'haz', 'crea', 'genera',
                           'para la', 'sobre', 'acerca de']
        
        tema = user_input.lower()
        for palabra in palabras_excluir:
            tema = tema.replace(palabra, '')
        
        # Limpiar espacios extras y caracteres especiales
        tema = re.sub(r'[^\w\s]', '', tema)
        tema = ' '.join(tema.split())
        
        return tema.strip() if tema.strip() else user_input
    except Exception:
        return user_input

def detectar_tipo_solicitud(user_input):
    """Detecta el tipo de solicitud del usuario"""
    input_lower = user_input.lower()
    
    if any(palabra in input_lower for palabra in ['planteamiento', 'problema', 'pregunta investigación']):
        return "planteamiento"
    elif any(palabra in input_lower for palabra in ['objetivos', 'metas', 'propósitos']):
        return "objetivos"
    elif any(palabra in input_lower for palabra in ['metodología', 'método', 'diseño', 'enfoque']):
        return "metodologia"
    elif any(palabra in input_lower for palabra in ['variables', 'operacional']):
        return "variables"
    elif any(palabra in input_lower for palabra in ['resumen', 'síntesis', 'analizar']):
        return "resumen"
    else:
        return "general"

# FUNCIONES DE GENERACIÓN MEJORADAS
def generar_planteamiento_estructurado(tema, contexto=""):
    """Genera un planteamiento del problema bien estructurado"""
    try:
        client = genai.Client()
        
        prompt = f"""
        Como experto en metodología de investigación, genera un PLANTEAMIENTO DEL PROBLEMA 
        académico y profesional sobre el tema: "{tema}"
        
        Contexto adicional: {contexto}
        
        Estructura tu respuesta en los siguientes componentes:
        
        # DESCRIPCIÓN DEL PROBLEMA
        [Describe claramente el problema de investigación, su contexto y magnitud]
        
        # JUSTIFICACIÓN
        [Explica la relevancia académica, práctica y social de investigar este problema]
        
        # DELIMITACIÓN
        [Especifica el alcance, población y contexto del estudio]
        
        # PREGUNTAS DE INVESTIGACIÓN
        [Formula 3-5 preguntas de investigación específicas y relevantes]
        
        Usa un lenguaje académico formal en español. Sé específico y evgeneralizaciones.
        """
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(temperature=0.3)
        )
        
        return response.text.strip()
        
    except Exception as e:
        return f"Error al generar planteamiento: {e}"

def generar_objetivos_estructurados(tema, contexto=""):
    """Genera objetivos de investigación estructurados"""
    try:
        client = genai.Client()
        
        prompt = f"""
        Para la investigación sobre: "{tema}"
        
        Contexto: {contexto}
        
        Genera objetivos de investigación que incluyan:
        
        OBJETIVO GENERAL:
        [Un objetivo principal amplio]
        
        OBJETIVOS ESPECÍFICOS:
        1. [Objetivo específico 1]
        2. [Objetivo específico 2] 
        3. [Objetivo específico 3]
        4. [Objetivo específico 4]
        
        Los objetivos deben ser SMART (específicos, medibles, alcanzables, relevantes, temporales)
        y coherentes con el tema de investigación.
        """
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(temperature=0.3)
        )
        
        return response.text.strip()
        
    except Exception as e:
        return f"Error al generar objetivos: {e}"

def generar_respuesta_general(tema, user_input):
    """Genera respuesta general del asistente"""
    try:
        client = genai.Client()
        
        prompt = f"""
        Eres un asistente de investigación académica especializado en {tema}.
        
        El usuario pregunta: "{user_input}"
        
        Proporciona una respuesta útil, académica y bien fundamentada. Si es apropiado, sugiere:
        - Enfoques metodológicos
        - Fuentes de datos relevantes
        - Conceptos clave para investigar
        - Posibles líneas de investigación
        
        Respuesta en español, formato claro y profesional.
        """
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(temperature=0.7)
        )
        
        return response.text.strip()
        
    except Exception as e:
        return f"Error en la respuesta: {e}"

# FUNCIÓN DE CHAT COMPLETAMENTE REDISEÑADA
def chat_con_agente(agente, user_input, contexto_usuario=""):
    """Función principal del chat completamente mejorada"""
    try:
        # Extraer el tema real de la consulta
        tema_real = extraer_tema_principal(user_input)
        tipo_solicitud = detectar_tipo_solicitud(user_input)
        
        # Generar respuesta según el tipo de solicitud
        if tipo_solicitud == "planteamiento":
            respuesta = generar_planteamiento_estructurado(tema_real, contexto_usuario)
            
        elif tipo_solicitud == "objetivos":
            respuesta = generar_objetivos_estructurados(tema_real, contexto_usuario)
            
        elif tipo_solicitud == "metodologia":
            respuesta = generar_sugerencias_metodologicas("", contexto_usuario or tema_real)
            
        elif tipo_solicitud == "variables":
            respuesta = sugerir_variables_operativas("", contexto_usuario or tema_real)
            
        elif tipo_solicitud == "resumen":
            if hasattr(st.session_state, 'resumen_actual'):
                respuesta = st.session_state.resumen_actual
            else:
                respuesta = "Primero realiza una búsqueda académica para generar resúmenes."
                
        else:
            respuesta = generar_respuesta_general(tema_real, user_input)
        
        # Guardar en memoria
        agente.agregar_memoria(user_input, respuesta)
        return respuesta
        
    except Exception as e:
        return f"Error en el chat: {e}"

# FUNCIONES ORIGINALES (MANTENIDAS)
def preparar_texto_para_gemini(articulos):
    texto = ""
    for art in articulos:
        texto += f"Título: {art['titulo']}\nAutores: {art['autor']}\nAño: {art['año']}\nURL: {art['url']}\n\n"
    return texto

def resumir_articulos_con_gemini(texto):
    try:
        client = genai.Client()
        prompt = (f"Analiza estos artículos académicos para resumir temas principales, "
                 f"vacíos de investigación y sugerir líneas de trabajo para un proyecto:\n\n{texto}")
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(temperature=0.2)
        )
        return response.text
    except Exception as e:
        return f"Error al generar resumen: {e}"

def generar_planteamiento_problema(texto_resumen, contexto_usuario):
    try:
        client = genai.Client()
        prompt = (
            f"Con base en la siguiente síntesis de literatura científica:\n{texto_resumen}\n\n"
            f"Considerando que el investigador se interesa en: {contexto_usuario}\n"
            "Genera un planteamiento del problema estructurado para un proyecto de investigación que incluya:\n"
            ". Descripción clara del problema\n"
            ". Justificación de la investigación\n"
            ". Delimitación del campo o población\n"
            ". Preguntas de investigación bien formuladas\n"
        )
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(temperature=0.3)
        )
        return response.text
    except Exception as e:
        return f"Error al generar planteamiento: {e}"

def generar_objetivos_investigacion(texto_resumen, contexto_usuario):
    try:
        client = genai.Client()
        prompt = (
            f"Con base en la siguiente síntesis de literatura científica y el contexto del investigador:\n{texto_resumen}\n\n"
            f"Considerando que el investigador desea enfocarse en: {contexto_usuario}\n"
            "Genera los objetivos de investigación para un proyecto académico que incluyan:\n"
            "- Un objetivo general\n"
            "- Tres a cinco objetivos específicos coherentes con el problema y la justificación\n"
        )
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(temperature=0.3)
        )
        return response.text
    except Exception as e:
        return f"Error al generar objetivos: {e}"

def sugerir_variables_operativas(texto_resumen, contexto_usuario):
    try:
        client = genai.Client()
        prompt = (
            f"Basado en esta síntesis de literatura científica:\n{texto_resumen}\n\n"
            f"Considerando que el investigador trabaja en el contexto: {contexto_usuario}\n"
            "Genera una propuesta de variables de investigación (dependientes, independientes, mediadoras, moderadoras) "
            "y sus definiciones operativas claras, específicas y coherentes con el problema y objetivos de investigación."
        )
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(temperature=0.3)
        )
        return response.text
    except Exception as e:
        return f"Error al generar variables: {e}"

def generar_sugerencias_metodologicas(texto_resumen, contexto_usuario):
    try:
        client = genai.Client()
        prompt = (
            f"Considerando la siguiente síntesis de literatura científica:\n{texto_resumen}\n\n"
            f"Y el contexto de investigación: {contexto_usuario}\n"
            "Genera sugerencias detalladas sobre:\n"
            "- Tipo o enfoque metodológico (cuantitativo, cualitativo, mixto)\n"
            "- Diseño de investigación apropiado\n"
            "- Técnicas e instrumentos de recolección de datos recomendadas\n"
            "- Estrategias para el análisis de datos\n"
            "Estas sugerencias deben ser coherentes con el planteamiento del problema y los objetivos de investigación."
        )
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(temperature=0.3)
        )
        return response.text
    except Exception as e:
        return f"Error al generar sugerencias metodológicas: {e}"

def generar_bibliografia(referencias, estilo="APA"):
    def formatear_cita(ref):
        autor = ref.get("autor", "Desconocido").strip()
        año = ref.get("año", "").strip() or "s.f."
        titulo = ref.get("titulo", "").strip()
        publicacion = ref.get("publicacion", "").strip()
        url = ref.get("url", "").strip()

        if estilo.upper() == "APA":
            partes = []
            partes.append(f"{autor} ({año})")
            if titulo:
                partes.append(f"{titulo}.")
            if publicacion:
                partes.append(f"{publicacion}.")
            cita = " ".join(partes)
            if url:
                cita += f" Recuperado de {url}"
        elif estilo.upper() == "UPEL":
            partes = []
            partes.append(f"{autor}. ({año}).")
            if titulo:
                partes.append(f"{titulo}.")
            if publicacion:
                partes.append(f"{publicacion}.")
            cita = " ".join(partes)
            if url:
                cita += f" Disponible en: {url}"
        else:
            cita = f"{autor}, {año}, {titulo}"
        return " ".join(cita.split())

    bibliografia = [formatear_cita(ref) for ref in referencias]
    return "\n\n".join(bibliografia)

# FUNCIÓN PRINCIPAL DE STREAMLIT
def main():
    # Configuración de página
    st.set_page_config(
        page_title="Asistente de Investigación Inteligente",
        page_icon="🔬",
        layout="wide"
    )
    
    # Header principal
    st.markdown("""
    <div class="main-header" style="text-align: center; padding: 20px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border-radius: 10px; margin-bottom: 20px;">
        <h1>🔬 Asistente de Investigación Inteligente para Postgrado</h1>
        <p>Chatbot mejorado con procesamiento avanzado de lenguaje natural</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.image("https://cdn-icons-png.flaticon.com/512/2913/2913517.png", width=80)
        st.title("Configuración")
        
        st.markdown("---")
        st.subheader("🔍 Búsqueda")
        max_results = st.slider("Resultados por base", 1, 10, 5)
        
        st.markdown("---")
        st.subheader("📤 Exportación")
        formato_exportacion = st.selectbox("Formato preferido", ["CSV", "DOCX", "TXT"])
        
        st.markdown("---")
        st.subheader("ℹ️ Información")
        st.info("""
        **Bases de datos gratuitas:**
        - Semantic Scholar
        - SciELO
        
        **Funcionalidades MEJORADAS:**
        - Chatbot inteligente con NLP
        - Detección automática de temas
        - Generación contextualizada
        - Exportación de resultados
        """)
    
    # Inicializar session state
    if 'articulos' not in st.session_state:
        st.session_state.articulos = []
    if 'agente' not in st.session_state:
        st.session_state.agente = AgenteConMemoria()
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'contexto_actual' not in st.session_state:
        st.session_state.contexto_actual = ""
    
    # Pestañas principales
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "🏠 Inicio", 
        "🔍 Búsqueda Académica", 
        "📊 Generar Proyecto", 
        "💬 Chat con Agente",
        "📅 Exportar Resultados"
    ])
    
    # Pestaña de Inicio
    with tab1:
        st.markdown("### 💡 Bienvenido al Asistente de Investigación Inteligente")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("""
            ### 🚀 ¿Qué puedes hacer con la VERSIÓN MEJORADA?
            
            **🔍 Chatbot Inteligente MEJORADO**
            - Procesamiento avanzado de lenguaje natural
            - Detección automática de temas de investigación
            - Generación contextualizada de planteamientos
            - Respuestas específicas y relevantes
            
            **📚 Búsqueda Académica Integrada**
            - Buscar en Semantic Scholar (Internacional)
            - Buscar en SciELO (América Latina)
            - Resultados en tiempo real
            - Filtrado por relevancia
            
            **🤖 Análisis con Inteligencia Artificial**
            - Resúmenes automáticos de literatura
            - Detección de tendencias de investigación
            - Identificación de vacíos en la literatura
            - Sugerencias de líneas de investigación
            
            **📋 Desarrollo de Proyectos**
            - Generar planteamientos de problema
            - Definir objetivos de investigación
            - Sugerir variables operativas
            - Proponer metodologías apropiadas
            """)
        
        with col2:
            st.markdown("""
            ### 🎯 Comenzar es fácil:
            
            1. **Ve a la pestaña 💬 Chat con Agente**
            2. **Escribe tu pregunta naturalmente**
            3. **El sistema detectará automáticamente tu tema**
            4. **Recibe respuestas contextualizadas**
            
            ### 💬 Ejemplos de preguntas:
            - *"Formula el planteamiento del problema sobre competencias digitales y resiliencia en entornos automatizados"*
            - *"Genera objetivos para investigación sobre inteligencia artificial en educación"*
            - *"Sugiere metodología para estudiar impacto de redes sociales"*
            
            ### 🗃️ Bases disponibles:
            ☑ Semantic Scholar  
            ☑ SciELO  
            ☑ Acceso gratuito  
            """)
        
        st.markdown("---")
        st.subheader("📊 Estadísticas de la sesión")
        col_stats1, col_stats2, col_stats3 = st.columns(3)
        with col_stats1:
            st.metric("Artículos encontrados", len(st.session_state.articulos))
        with col_stats2:
            st.metric("Consultas al agente", len(st.session_state.agente.obtener_memoria()))
        with col_stats3:
            st.metric("Mensajes en chat", len(st.session_state.chat_history))
    
    # Pestaña de Búsqueda Académica
    with tab2:
        st.markdown("### 🔍 Búsqueda en Bases Académicas")
        
        col_search1, col_search2 = st.columns([2, 1])
        
        with col_search1:
            consulta = st.text_input(
                "🔍 Tema de investigación:",
                placeholder="Ej: inteligencia artificial educación secundaria",
                help="Describe tu tema de investigación con palabras clave específicas"
            )
        
        with col_search2:
            contexto = st.text_input(
                "🎯 Contexto específico:",
                placeholder="Ej: análisis de impacto en Venezuela",
                help="Especifica el contexto o enfoque particular"
            )
        
        # Selección de bases de datos
        st.subheader("Selecciona las bases de datos:")
        col_db1, col_db2 = st.columns(2)
        
        with col_db1:
            usar_semantic = st.checkbox("Semantic Scholar", value=True, help="Artículos en inglés e internacionales")
        with col_db2:
            usar_scielo = st.checkbox("SciELO", value=True, help="Revistas científicas de América Latina")
        
        if st.button("🔍 Ejecutar Búsqueda Integral", type="primary", use_container_width=True):
            if not consulta:
                st.error("📌 Por favor ingresa un tema de investigación")
            else:
                with st.spinner("🔍 Buscando en bases académicas..."):
                    # Ejecutar búsquedas
                    articulos_encontrados = []
                    
                    if usar_semantic:
                        with st.expander("🌐 Semantic Scholar", expanded=True):
                            resultados_ss = buscar_semantic_scholar(consulta, max_results)
                            articulos_encontrados.extend(resultados_ss)
                            st.success(f"📄 Encontrados: {len(resultados_ss)} artículos")
                    
                    if usar_scielo:
                        with st.expander("🌎 SciELO", expanded=True):
                            resultados_scielo = buscar_scielo(consulta, max_results)
                            articulos_encontrados.extend(resultados_scielo)
                            st.success(f"📄 Encontrados: {len(resultados_scielo)} artículos")
                    
                    # Guardar resultados
                    st.session_state.articulos = articulos_encontrados
                    st.session_state.consulta_actual = consulta
                    st.session_state.contexto_actual = contexto
                    
                    if articulos_encontrados:
                        st.success(f"✅ Búsqueda completada! Se encontraron {len(articulos_encontrados)} artículos")
                        
                        # Generar y mostrar resumen con Gemini
                        with st.expander("🤖 Resumen y Análisis con IA", expanded=True):
                            texto_gemini = preparar_texto_para_gemini(articulos_encontrados)
                            if texto_gemini.strip():
                                resumen = resumir_articulos_con_gemini(texto_gemini)
                                st.markdown(resumen)
                                st.session_state.resumen_actual = resumen
                            else:
                                st.warning("No se encontraron artículos. Intenta con otros términos de búsqueda.")
                    else:
                        st.warning("No se encontraron artículos con los criterios especificados.")
        
        # Mostrar resultados detallados
        if st.session_state.articulos:
            st.markdown("---")
            st.subheader("📚 Artículos Encontrados")
            
            for i, articulo in enumerate(st.session_state.articulos):
                with st.expander(f"📖 {articulo['titulo']}", key=f"art_{i}"):
                    col_info1, col_info2 = st.columns([3, 1])
                    
                    with col_info1:
                        st.write(f"**Autores:** {articulo['autor']}")
                        st.write(f"**Año:** {articulo['año']}")
                        st.write(f"**Publicación:** {articulo['publicacion']}")
                        st.write(f"**Fuente:** {articulo['fuente']}")
                    
                    with col_info2:
                        if articulo['url']:
                            st.markdown(f"[🔗 Enlace al artículo]({articulo['url']})")
                        else:
                            st.write("🔗 Enlace no disponible")
    
    # Pestaña de Generar Proyecto
    with tab3:
        st.markdown("### 📊 Generar Elementos de Investigación")
        
        if not st.session_state.articulos:
            st.warning("⚠️ Primero realiza una búsqueda en la pestaña anterior para generar contenido.")
        else:
            col_gen1, col_gen2 = st.columns(2)
            
            with col_gen1:
                if st.button("📝 Generar Planteamiento del Problema", use_container_width=True):
                    with st.spinner("Generando planteamiento del problema..."):
                        if hasattr(st.session_state, 'resumen_actual'):
                            planteamiento = generar_planteamiento_problema(
                                st.session_state.resumen_actual,
                                st.session_state.contexto_actual
                            )
                            st.session_state.planteamiento = planteamiento
                            st.markdown(planteamiento)
                        else:
                            st.error("Primero genera un resumen en la pestaña de Búsqueda")
            
            with col_gen2:
                if st.button("🎯 Generar Objetivos de Investigación", use_container_width=True):
                    with st.spinner("Generando objetivos de investigación..."):
                        if hasattr(st.session_state, 'resumen_actual'):
                            objetivos = generar_objetivos_investigacion(
                                st.session_state.resumen_actual,
                                st.session_state.contexto_actual
                            )
                            st.session_state.objetivos = objetivos
                            st.markdown(objetivos)
                        else:
                            st.error("Primero genera un resumen en la pestaña de Búsqueda")
            
            st.markdown("---")
            col_gen3, col_gen4 = st.columns(2)
            
            with col_gen3:
                if st.button("📊 Sugerir Variables Operativas", use_container_width=True):
                    with st.spinner("Generando variables operativas..."):
                        if hasattr(st.session_state, 'resumen_actual'):
                            variables = sugerir_variables_operativas(
                                st.session_state.resumen_actual,
                                st.session_state.contexto_actual
                            )
                            st.session_state.variables = variables
                            st.markdown(variables)
                        else:
                            st.error("Primero genera un resumen en la pestaña de Búsqueda")
            
            with col_gen4:
                if st.button("🔬 Generar Sugerencias Metodológicas", use_container_width=True):
                    with st.spinner("Generando sugerencias metodológicas..."):
                        if hasattr(st.session_state, 'resumen_actual'):
                            metodologia = generar_sugerencias_metodologicas(
                                st.session_state.resumen_actual,
                                st.session_state.contexto_actual
                            )
                            st.session_state.metodologia = metodologia
                            st.markdown(metodologia)
                        else:
                            st.error("Primero genera un resumen en la pestaña de Búsqueda")
    
    # Pestaña de Chat con Agente - COMPLETAMENTE REDISEÑADA
    with tab4:
        st.markdown("### 💬 Chat Inteligente con el Agente de Investigación")
        
        st.info("""
        🎯 **Pregunta naturalmente sobre:**
        - Planteamientos de problemas de investigación
        - Objetivos y metodologías
        - Análisis de datos y variables
        - Redacción académica y referencias
        - Diseño de instrumentos de investigación
        
        💡 **Ejemplo:** *"Formula el planteamiento del problema sobre competencias digitales en entornos automatizados"*
        """)
        
        # Configuración de contexto
        st.session_state.contexto_actual = st.text_input(
            "🎯 Contexto de investigación (opcional):",
            placeholder="Ej: educación superior en Latinoamérica",
            help="Proporciona contexto para respuestas más precisas"
        )
        
        # Mostrar historial de chat
        for mensaje in st.session_state.chat_history:
            with st.chat_message(mensaje["role"]):
                st.markdown(mensaje["content"])
        
        # Input de chat
        if prompt := st.chat_input("Escribe tu pregunta sobre investigación..."):
            # Agregar mensaje del usuario
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Respuesta del agente MEJORADO
            with st.chat_message("assistant"):
                with st.spinner("🤔 Analizando tu consulta..."):
                    # Usar la función de chat mejorada
                    respuesta = chat_con_agente(
                        st.session_state.agente, 
                        prompt, 
                        st.session_state.contexto_actual
                    )
                    st.markdown(respuesta)
            
            # Agregar respuesta al historial
            st.session_state.chat_history.append({"role": "assistant", "content": respuesta})
        
        # Botones de acción
        col_chat1, col_chat2, col_chat3 = st.columns(3)
        with col_chat1:
            if st.button("🗑️ Limpiar Conversación", use_container_width=True):
                st.session_state.chat_history = []
                st.session_state.agente.limpiar_memoria()
                st.rerun()
        
        with col_chat2:
            if st.button("💡 Ejemplo de Planteamiento", use_container_width=True):
                ejemplo = "Formula el planteamiento del problema sobre competencias digitales para la resiliencia en entornos automatizados con IA"
                st.session_state.chat_history.append({"role": "user", "content": ejemplo})
                with st.chat_message("assistant"):
                    with st.spinner("🤔 Generando ejemplo..."):
                        respuesta = chat_con_agente(
                            st.session_state.agente, 
                            ejemplo, 
                            "entornos laborales digitalizados"
                        )
                        st.markdown(respuesta)
                st.session_state.chat_history.append({"role": "assistant", "content": respuesta})
                st.rerun()
        
        with col_chat3:
            if st.button("📊 Ver Estadísticas", use_container_width=True):
                memoria = st.session_state.agente.obtener_memoria()
                st.info(f"**Estadísticas del Agente:**\n- Consultas procesadas: {len(memoria)}\n- Tema actual: {st.session_state.contexto_actual}")
    
    # Pestaña de Exportar Resultados
    with tab5:
        st.markdown("### 📤 Exportar Resultados de Investigación")
        
        if not st.session_state.articulos:
            st.warning("📌 No hay resultados para exportar. Realiza primero una búsqueda.")
        else:
            col_exp1, col_exp2, col_exp3 = st.columns(3)
            
            with col_exp1:
                st.subheader("📚 Artículos Encontrados")
                
                # Exportar CSV
                campos_csv = ["titulo", "autor", "año", "publicacion", "fuente", "url"]
                articulos_exportar = []
                
                for art in st.session_state.articulos:
                    articulos_exportar.append({
                        "titulo": art.get("titulo", ""),
                        "autor": art.get("autor", ""),
                        "año": art.get("año", "s.f."),
                        "publicacion": art.get("publicacion", ""),
                        "fuente": art.get("fuente", ""),
                        "url": art.get("url", "")
                    })
                
                # Crear CSV en memoria
                output = io.StringIO()
                writer = csv.DictWriter(output, fieldnames=campos_csv)
                writer.writeheader()
                for item in articulos_exportar:
                    writer.writerow(item)
                csv_data = output.getvalue()
                
                st.download_button(
                    label="📥 Descargar Artículos (CSV)",
                    data=csv_data,
                    file_name=f"articulos_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            
            with col_exp2:
                st.subheader("📄 Documentos Generados")
                
                if 'planteamiento' in st.session_state:
                    # Crear DOCX en memoria
                    doc = Document()
                    doc.add_heading("Planteamiento del Problema", level=1)
                    doc.add_paragraph(st.session_state.planteamiento)
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    st.download_button(
                        label="📄 Planteamiento (DOCX)",
                        data=buffer,
                        file_name="planteamiento_problema.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                if 'objetivos' in st.session_state:
                    # Exportar objetivos como TXT
                    objetivos_txt = st.session_state.objetivos
                    st.download_button(
                        label="🎯 Objetivos (TXT)",
                        data=objetivos_txt,
                        file_name="objetivos_investigacion.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
            
            with col_exp3:
                st.subheader("📖 Bibliografía")
                
                # Generar bibliografía
                referencias = []
                for art in st.session_state.articulos:
                    ref = {
                        "autor": art["autor"],
                        "año": art["año"] if art["año"] else "s.f.",
                        "titulo": art["titulo"],
                        "publicacion": art["publicacion"],
                        "url": art["url"],
                    }
                    referencias.append(ref)
                
                bibliografia_apa = generar_bibliografia(referencias, estilo="APA")
                
                st.download_button(
                    label="📚 Bibliografia APA (TXT)",
                    data=bibliografia_apa,
                    file_name="bibliografia_APA.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            st.markdown("---")
            st.subheader("📊 Estadísticas de Exportación")
            
            col_stats1, col_stats2, col_stats3, col_stats4 = st.columns(4)
            
            with col_stats1:
                st.metric("Artículos", len(st.session_state.articulos))
            
            with col_stats2:
                fuentes = set(art['fuente'] for art in st.session_state.articulos)
                st.metric("Fuentes", len(fuentes))
            
            with col_stats3:
                años = [art['año'] for art in st.session_state.articulos if art['año'].isdigit()]
                if años:
                    st.metric("Rango de años", f"{min(años)}-{max(años)}")
                else:
                    st.metric("Rango de años", "N/A")
            
            with col_stats4:
                st.metric("Formatos", "CSV, DOCX, TXT")

if __name__ == "__main__":
    main()
